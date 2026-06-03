"""
doc_builder.py — python-docx Manuscript Builder
===============================================
Manages a single master manuscript.docx that grows via append-only
operations. Never rewrites the whole document — always opens and saves.

Formatting rules:
  - Normal prose clips are appended as runs inside the CURRENT active
    paragraph (space-concatenated, not newline-separated).
  - /next_paragraph   -> seal current paragraph, start a fresh one.
  - /next_chapter     -> seal paragraph, page break, bold Heading 1 title.
  - First cold start  -> creates a document with a title page run.

Marathi typography:
  - All body runs use 'Shobhika' with 'Tiro Marathi' as a fallback
    theme font, ensuring correct complex-script ligature rendering
    during LibreOffice headless PDF compilation.

Atomic save: writes to manuscript.docx.tmp then os.replace() to prevent
corruption if a new audio clip arrives while a conversion is in progress.

Python 3.8+ compatible (no X | Y union type-hint syntax).
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import MANUSCRIPT_DOCX, BOOK_DIR

# ── Preferred Marathi fonts (in priority order) ───────────────────────────────
# Shobhika is a high-quality open-source Devanagari font; Tiro Marathi is the
# Google Fonts alternative. Both render Marathi conjuncts correctly in LibreOffice.
_BODY_FONT_PRIMARY  = "Shobhika"
_BODY_FONT_FALLBACK = "Tiro Marathi"
_BODY_FONT_SIZE_PT  = 12

_HEADING_FONT       = "Shobhika"
_HEADING_FONT_SIZE  = 18

# ── Internal state ────────────────────────────────────────────────────────────
# We cache the docx Document object to avoid re-parsing on every clip.
# Optional[Document] used instead of Document | None for Python 3.8 compat.

_doc = None                  # type: Optional[Document]
_active_paragraph = None     # current docx paragraph we are appending runs to


def _load_or_create():
    # type: () -> Document
    """
    Open existing manuscript.docx or create a fresh document.
    Sets _active_paragraph to the last body paragraph.
    """
    global _doc, _active_paragraph

    if MANUSCRIPT_DOCX.exists():
        _doc = Document(str(MANUSCRIPT_DOCX))
        print("[DocBuilder] Opened existing manuscript ({} paragraphs).".format(
            len(_doc.paragraphs)
        ))
        # Resume appending from the last paragraph in the document
        if _doc.paragraphs:
            _active_paragraph = _doc.paragraphs[-1]
        else:
            _active_paragraph = _doc.add_paragraph()
    else:
        _doc = Document()
        _apply_document_defaults(_doc)
        _create_title_page(_doc)
        _active_paragraph = _doc.add_paragraph()
        _active_paragraph.style = _doc.styles["Normal"]
        print("[DocBuilder] Created fresh manuscript.docx.")

    return _doc


def _apply_document_defaults(doc):
    # type: (Document) -> None
    """Set Marathi font, margins, and line-spacing defaults for the whole document."""
    style = doc.styles["Normal"]
    font = style.font
    # Primary Marathi font — LibreOffice will use this when rendering
    font.name = _BODY_FONT_PRIMARY
    font.size = Pt(_BODY_FONT_SIZE_PT)

    # Set the complex-script (cs) font so Devanagari glyphs pick up correctly
    _set_complex_script_font(style._element, _BODY_FONT_PRIMARY)

    style.paragraph_format.line_spacing = Pt(20)
    style.paragraph_format.space_after  = Pt(8)

    # Page margins (book-style)
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)


def _set_complex_script_font(rPr_parent, font_name):
    # type: (object, str) -> None
    """
    Inject <w:rFonts> with cs= and eastAsia= attributes so that LibreOffice
    uses the Devanagari font for complex-script rendering.
    Works on both paragraph style rPr and individual run rPr elements.
    """
    try:
        # Navigate to the rPr element within the style's pPr or directly
        # For style elements we reach rPr via the style XML
        rPr = rPr_parent.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            rPr_parent.append(rPr)

        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)

        rFonts.set(qn("w:ascii"),     font_name)
        rFonts.set(qn("w:hAnsi"),     font_name)
        rFonts.set(qn("w:cs"),        font_name)   # complex-script (Devanagari)
        rFonts.set(qn("w:eastAsia"),  font_name)
    except Exception as exc:
        print("[DocBuilder] _set_complex_script_font warning: {}".format(exc))


def _apply_marathi_font_to_run(run):
    """Force Shobhika on a run's ascii, hAnsi, and complex-script slots."""
    run.font.name = _BODY_FONT_PRIMARY
    run.font.size = Pt(_BODY_FONT_SIZE_PT)
    # Inject cs font via XML so LibreOffice picks it up
    _set_complex_script_font(run._r, _BODY_FONT_PRIMARY)


def _create_title_page(doc):
    # type: (Document) -> None
    """Add a minimal title block at the top of a fresh document."""
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("माझे पुस्तक")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = _BODY_FONT_PRIMARY
    _set_complex_script_font(run._r, _BODY_FONT_PRIMARY)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        "सुरुवात: {}".format(datetime.now().strftime("%d %B %Y"))
    )
    sub_run.font.size = Pt(11)
    sub_run.font.name = _BODY_FONT_PRIMARY
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _set_complex_script_font(sub_run._r, _BODY_FONT_PRIMARY)

    doc.add_paragraph()  # blank spacer


def _get_doc():
    # type: () -> Document
    """Return cached document, loading from disk if not yet initialised."""
    global _doc
    if _doc is None:
        _load_or_create()
    return _doc


def _ensure_active_paragraph():
    # type: () -> None
    """Guarantee _active_paragraph points to a valid paragraph."""
    global _active_paragraph
    doc = _get_doc()
    if _active_paragraph is None:
        _active_paragraph = doc.add_paragraph()
        _active_paragraph.style = doc.styles["Normal"]


# ── Public API ────────────────────────────────────────────────────────────────

def initialise():
    # type: () -> None
    """Call once at engine startup to load or create the manuscript."""
    _load_or_create()


def append_text(polished_text):
    # type: (str) -> None
    """
    Append polished prose to the current active paragraph.
    Multiple clips are joined with a space — NOT newlines.
    All runs are forced to Shobhika for correct Devanagari rendering.
    """
    if not polished_text or not polished_text.strip():
        return

    _ensure_active_paragraph()

    # Add a space joiner if the paragraph already has content
    existing_text = _active_paragraph.text
    if existing_text.strip():
        space_run = _active_paragraph.add_run(" ")
        _apply_marathi_font_to_run(space_run)

    text_run = _active_paragraph.add_run(polished_text.strip())
    _apply_marathi_font_to_run(text_run)
    _save()


def flush_paragraph():
    # type: () -> None
    """
    Seal the current paragraph and start a new empty one.
    Triggered by /next_paragraph or 'pudil paricched'.
    """
    global _active_paragraph
    doc = _get_doc()

    _active_paragraph = doc.add_paragraph()
    _active_paragraph.style = doc.styles["Normal"]
    _save()
    print("[DocBuilder] New paragraph started.")


def flush_chapter(title):
    # type: (str) -> None
    """
    Insert a page break then a bold Heading 1 chapter title (Shobhika font),
    and reset the active paragraph to a fresh body paragraph.
    Triggered by /next_chapter [Title] or 'pudil dhada [Title]'.
    """
    global _active_paragraph
    doc = _get_doc()

    # Page break via OxmlElement — most reliable approach in python-docx
    page_break_para = doc.add_paragraph()
    _insert_page_break(page_break_para)

    # Heading 1 with explicit bold run in Shobhika
    heading_para = doc.add_heading(level=1)
    heading_run = heading_para.add_run(title.strip())
    heading_run.bold = True
    heading_run.font.name = _HEADING_FONT
    heading_run.font.size = Pt(_HEADING_FONT_SIZE)
    _set_complex_script_font(heading_run._r, _HEADING_FONT)

    # Fresh body paragraph after the heading
    _active_paragraph = doc.add_paragraph()
    _active_paragraph.style = doc.styles["Normal"]

    _save()
    print('[DocBuilder] New chapter: "{}"'.format(title.strip()))


def _insert_page_break(paragraph):
    """Insert a manual page break run into a paragraph."""
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _save():
    # type: () -> None
    """
    Atomically save the document: write to .tmp then os.replace().
    Guarantees manuscript.docx is never left in a half-written state
    even if a new audio clip arrives mid-write.
    """
    global _doc
    tmp_path = BOOK_DIR / "manuscript.docx.tmp"
    _doc.save(str(tmp_path))
    os.replace(str(tmp_path), str(MANUSCRIPT_DOCX))


def reload():
    # type: () -> None
    """Force reload from disk (useful after external edits or crash recovery)."""
    global _doc, _active_paragraph
    _doc = None
    _active_paragraph = None
    _load_or_create()
